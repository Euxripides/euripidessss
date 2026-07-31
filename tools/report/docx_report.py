#!/usr/bin/env python3
"""案件分析报告 DOCX 生成器。

用法: python docx_report.py <case-report.json> <output.docx>

格式要求:
- 字体: 仿宋 (FangSong), 小四 (12pt)
- 不出现多余标题样式和横线: 结构段落用加粗区分, 不使用 Heading 样式
  (Heading 样式会带主题色/边框线), 不插入分隔线。
"""
import json
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size_pt=12, bold=False):
    run.font.name = "Times New Roman"  # ASCII 字体
    run.font.size = Pt(size_pt)
    run.bold = bold
    # 东亚字体: 仿宋
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "仿宋")


def add_para(doc, text, bold=False, size_pt=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold)
    return p


def main():
    if len(sys.argv) != 3:
        print("usage: docx_report.py <case-report.json> <output.docx>")
        return 2
    with open(sys.argv[1], encoding="utf-8") as f:
        case = json.load(f)

    doc = Document()
    # Normal 样式默认仿宋小四
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    add_para(doc, "案件分析报告", bold=True, size_pt=16, space_after=12)
    add_para(doc, f"案件编号：{case.get('case_id', '')}")
    add_para(doc, f"调查员：{case.get('investigator', '')}")
    add_para(doc, f"数据版本：{case.get('dataset_version', '')}")
    add_para(doc, f"状态：{case.get('status', '')}")
    add_para(doc, f"创建时间：{case.get('created_at', '')}")

    add_para(doc, "目标地址", bold=True)
    for addr in case.get("target_addresses", []):
        add_para(doc, f"· {addr}")

    summaries = case.get("summaries", {})
    if summaries:
        add_para(doc, "目标地址分析", bold=True)
        for addr, s in summaries.items():
            profile = s.get("profile", {})
            risk = s.get("risk", {})
            add_para(doc,
                     f"地址 {addr}：类型 {s.get('address_type', '')}，"
                     f"交易 {profile.get('transaction_count', 0)} 笔，"
                     f"流入 {s.get('in_count', 0)} / 流出 {s.get('out_count', 0)}，"
                     f"Token {profile.get('token_count', 0)} 个，"
                     f"风险 {risk.get('risk_score', 0)}（{risk.get('risk_level', '')}）")

    paths = case.get("trace_paths", [])
    if paths:
        add_para(doc, "资金流向", bold=True)
        add_para(doc, f"追踪路径 {len(paths)} 条")
        for p in paths[:20]:
            nodes = " → ".join(p.get("nodes", []))
            add_para(doc, f"· {nodes}（{p.get('hops', 0)} 跳）")

    common = case.get("common_sources", [])
    if common:
        add_para(doc, "公共资金来源", bold=True)
        for s in common:
            add_para(doc, f"· {s.get('address', '')}（覆盖 {s.get('count', 0)} 个目标）")

    sinks = case.get("common_sinks", [])
    if sinks:
        add_para(doc, "公共去向", bold=True)
        for s in sinks:
            add_para(doc, f"· {s.get('address', '')}（覆盖 {s.get('count', 0)} 个目标）")

    related = case.get("related_addresses", [])
    if related:
        add_para(doc, "关联地址", bold=True)
        for r in related[:10]:
            add_para(doc, f"· {r.get('address', '')}（score {r.get('shared_counterparty_score', 0):.3f}）")

    risks = case.get("risk_evidence", {})
    if risks:
        add_para(doc, "风险判断", bold=True)
        for addr, r in risks.items():
            add_para(doc, f"地址 {addr}：风险 {r.get('risk', {}).get('risk_score', 0)}"
                          f"（{r.get('risk', {}).get('risk_level', '')}），模式：{r.get('pattern', '')}")

    add_para(doc, "结论", bold=True)
    if case.get("status") == "COMPLETED":
        add_para(doc, "已完成调查分析，证据完整，详见证据文件（evidence.json / graph.json / timeline.csv）。")
    else:
        add_para(doc, f"调查未完成：{case.get('error', '')}")

    doc.save(sys.argv[2])
    print(f"docx saved: {sys.argv[2]}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
